#!/usr/bin/env python3
"""
Markdown to DOCX Converter Utility
----------------------------------
Converts Markdown transcripts into formatted Microsoft Word (.docx) documents,
supporting GitHub alerts/callout blocks, tables, titles, bold/italic text, and lists.
"""

import os
import sys
import re
import argparse

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("[!] python-docx is required. Install via: pip install python-docx")

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def convert_md_to_docx(md_path, output_docx_path=None):
    if not os.path.exists(md_path):
        print(f"[Error] Markdown file not found: {md_path}")
        return None
        
    if not output_docx_path:
        output_docx_path = os.path.splitext(md_path)[0] + ".docx"
        
    doc = Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    in_table = False
    table_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\r\n')
        
        # Table detection
        if line.strip().startswith("|"):
            in_table = True
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            in_table = False
            _process_table(doc, table_lines)
            table_lines = []
            
        # Headers
        if line.startswith("# "):
            h = doc.add_heading(level=1)
            _add_formatted_text(h, line[2:])
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
        elif line.startswith("## "):
            h = doc.add_heading(level=2)
            _add_formatted_text(h, line[3:])
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
        elif line.startswith("### "):
            h = doc.add_heading(level=3)
            _add_formatted_text(h, line[4:])
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(2)
        elif line.startswith("> [!"):
            # Callout block
            alert_type = re.match(r'> \[!(NOTE|IMPORTANT|WARNING|CAUTION|TIP)\]', line)
            alert_name = alert_type.group(1) if alert_type else "NOTE"
            alert_text = []
            i += 1
            while i < len(lines) and lines[i].startswith(">"):
                alert_text.append(lines[i].lstrip("> ").rstrip('\r\n'))
                i += 1
            i -= 1
            _add_callout_box(doc, alert_name, "\n".join(alert_text))
        elif line.startswith("* ") or line.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, line[2:])
            p.paragraph_format.space_after = Pt(2)
        elif line.strip():
            p = doc.add_paragraph()
            _add_formatted_text(p, line)
            p.paragraph_format.space_after = Pt(4)
            
        i += 1
        
    if in_table and table_lines:
        _process_table(doc, table_lines)
        
    doc.save(output_docx_path)
    print(f"[+] DOCX file created successfully: {output_docx_path}")
    return output_docx_path

def _add_formatted_text(paragraph, text):
    # Regex for bold **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Sub-split for code `text`
            subparts = re.split(r'(`.*?`)', part)
            for subpart in subparts:
                if subpart.startswith("`") and subpart.endswith("`"):
                    run = paragraph.add_run(subpart[1:-1])
                    run.font.name = 'Consolas'
                    run.font.color.rgb = RGBColor(180, 40, 40)
                else:
                    paragraph.add_run(subpart)

def _add_callout_box(doc, alert_type, text):
    colors = {
        "NOTE": "F0F7FF",
        "IMPORTANT": "F4EEFF",
        "WARNING": "FFF8E7",
        "CAUTION": "FFF0F0",
        "TIP": "EBFBF3"
    }
    fill_hex = colors.get(alert_type, "F9F9F9")
    
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, fill_hex)
    
    p = cell.paragraphs[0]
    title_run = p.add_run(f"[{alert_type}] ")
    title_run.bold = True
    _add_formatted_text(p, text)

def _process_table(doc, table_lines):
    rows = []
    for line in table_lines:
        if re.match(r'\|?\s*:?-+:?\s*\|', line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        if cells:
            rows.append(cells)
            
    if not rows:
        return
        
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    
    for r_idx, row in enumerate(rows):
        for c_idx, cell_value in enumerate(row):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                p = cell.paragraphs[0]
                _add_formatted_text(p, cell_value)
                if r_idx == 0:
                    set_cell_background(cell, "ECECEC")
                    for run in p.runs:
                        run.bold = True

def main():
    parser = argparse.ArgumentParser(description="Markdown to DOCX Converter")
    parser.add_argument("--input", required=True, help="Input Markdown file")
    parser.add_argument("--output", help="Output DOCX file")
    
    args = parser.parse_args()
    convert_md_to_docx(args.input, args.output)

if __name__ == "__main__":
    main()
